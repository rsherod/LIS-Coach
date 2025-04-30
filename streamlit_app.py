# Low-Intensity Strategies Coach (R. Sherod, fall 2024)
import streamlit as st
import google.generativeai as genai
from PIL import Image
import json
import os
from datetime import datetime
from io import BytesIO
import firebase_admin
from firebase_admin import credentials
from firebase_admin import firestore

# Import file format libraries (will be used in helper functions)
# Note: You may need to add these to your requirements.txt
# python-docx, reportlab

# Streamlit configuration
st.set_page_config(page_title="Streamlit Chatbot", layout="wide")

# Global CSS for other elements remains unchanged (if any)

# Initialize session state variables
if "form_submitted" not in st.session_state:
    st.session_state.form_submitted = False
if "form_responses" not in st.session_state:
    st.session_state.form_responses = {}
if "should_generate_response" not in st.session_state:
    st.session_state.should_generate_response = False
if "messages" not in st.session_state:
    st.session_state.messages = []
if "model_name" not in st.session_state:
    st.session_state.model_name = "gemini-2.0-flash"
if "temperature" not in st.session_state:
    st.session_state.temperature = 0.5
if "debug" not in st.session_state:
    st.session_state.debug = []
if "pdf_content" not in st.session_state:
    st.session_state.pdf_content = ""
if "chat_session" not in st.session_state:
    st.session_state.chat_session = None
if "uploaded_file" not in st.session_state:
    st.session_state.uploaded_file = None
if "active_strategy" not in st.session_state:
    st.session_state.active_strategy = None

# Helper function to load text files
def load_text_file(file_path):
    try:
        with open(file_path, 'r') as file:
            return file.read()
    except Exception as e:
        st.error(f"Error loading text file: {e}")
        return ""

# Helper function to load JSON files
def load_json_file(file_path):
    try:
        with open(file_path, 'r') as file:
            data = json.load(file)
            # If data is a list (array), convert it to a dictionary with strategy names as keys
            if isinstance(data, list):
                return {item["Strategy"]: item for item in data}
            return data
    except Exception as e:
        st.error(f"Error loading JSON file: {e}")
        return {}

# Helper functions for different file format exports
def get_chat_text_markdown():
    """Convert the chat messages to a downloadable markdown format"""
    chat_text = "# Low-Intensity Strategies Coach Chat Log\n\n"
    
    # Add timestamp
    chat_text += f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
    
    # Add strategy info if applicable
    if st.session_state.active_strategy:
        chat_text += f"Focus Strategy: {st.session_state.active_strategy}\n\n"
    
    # Add the messages
    for msg in st.session_state.messages:
        role = "Teacher" if msg["role"] == "user" else "Assistant"
        chat_text += f"## {role}:\n{msg['content']}\n\n"
    
    return chat_text

def get_chat_pdf():
    """Convert the chat messages to a PDF file using reportlab"""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from io import BytesIO
    
    # Create a buffer for the PDF
    buffer = BytesIO()
    
    # Create the PDF document
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Create a list to hold the elements
    elements = []
    
    # Add title
    title_style = styles["Title"]
    elements.append(Paragraph("Low-Intensity Strategies Coach Chat Log", title_style))
    elements.append(Spacer(1, 12))
    
    # Add timestamp
    normal_style = styles["Normal"]
    elements.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
    elements.append(Spacer(1, 12))
    
    # Add strategy info if applicable
    if st.session_state.active_strategy:
        elements.append(Paragraph(f"Focus Strategy: {st.session_state.active_strategy}", normal_style))
        elements.append(Spacer(1, 12))
    
    # Define styles for user and assistant
    user_style = ParagraphStyle(
        "UserStyle", 
        parent=styles["Heading2"],
        textColor=colors.darkblue
    )
    
    assistant_style = ParagraphStyle(
        "AssistantStyle", 
        parent=styles["Heading2"],
        textColor=colors.darkgreen
    )
    
    # Add the messages
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            role = "Teacher"
            heading_style = user_style
        else:
            role = "Assistant"
            heading_style = assistant_style
        
        # Add role heading
        elements.append(Paragraph(f"{role}:", heading_style))
        elements.append(Spacer(1, 6))
        
        # Add message content with proper line breaks
        text = msg["content"].replace('\n', '<br/>')
        elements.append(Paragraph(text, normal_style))
        elements.append(Spacer(1, 12))
    
    # Build the PDF
    doc.build(elements)
    
    # Get the value from the buffer
    pdf_value = buffer.getvalue()
    buffer.close()
    
    return pdf_value

def get_chat_docx():
    """Convert the chat messages to a Word document"""
    from docx import Document
    from io import BytesIO
    
    # Create document
    doc = Document()
    
    # Add title
    doc.add_heading('Low-Intensity Strategies Coach Chat Log', 0)
    
    # Add timestamp
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Add strategy info if applicable
    if st.session_state.active_strategy:
        doc.add_paragraph(f"Focus Strategy: {st.session_state.active_strategy}")
    
    # Add the messages
    for msg in st.session_state.messages:
        role = "Teacher" if msg["role"] == "user" else "Assistant"
        
        # Add role as heading
        doc.add_heading(f"{role}:", 2)
        
        # Add message content
        doc.add_paragraph(msg["content"])
    
    # Save to BytesIO object
    bytesio = BytesIO()
    doc.save(bytesio)
    bytesio.seek(0)
    
    # Return bytes
    return bytesio.getvalue()

# --- Add this function to save messages to Firestore ---
def save_message(user_input, bot_response):
    # Ensure db is initialized before trying to save
    if 'db' in globals() and db:
        try:
            # Data structure for a chat message
            message_data = {
                "user_input": user_input,
                "bot_response": bot_response,
                "timestamp": datetime.now(), # Add a server-side timestamp
                # Optional: Add active strategy for context
                "strategy": st.session_state.active_strategy if st.session_state.active_strategy else "main"
            }

            # Add a new document to the 'chats' collection in Firestore
            # Firestore will automatically generate a unique document ID
            doc_ref = db.collection("chats").add(message_data)

            # Optional: Add a debug message (be mindful of rate limits if too chatty)
            # st.session_state.debug.append(f"Message saved with ID: {doc_ref[1].id}")

        except Exception as e:
            st.session_state.debug.append(f"Firestore Save Error: {e}") # Log the error in debug
            st.error(f"Error saving message history: {e}") # Display error to user

    else:
        # This case happens if Firebase initialization failed
        st.session_state.debug.append("Firestore database not initialized. Cannot save message.")
        # st.warning("Message history cannot be saved (database not initialized).") # Optional user warning

# Load system instructions and strategy data
system_instructions = load_text_file('instructions.txt')
strategies_data = {}

# Define the paths to the JSON files
strategies_json_path = 'Strategies.json'

# Load the JSON file if it exists
if os.path.exists(strategies_json_path):
    strategies_data.update(load_json_file(strategies_json_path))
    st.session_state.debug.append(f"Loaded strategies from {strategies_json_path}")
else:
    st.session_state.debug.append(f"Warning: {strategies_json_path} not found")

# Function to build the complete system prompt
def build_system_prompt(active_strategy=None):
    # Start with base instructions
    prompt = system_instructions
    # Add strategy information
    if strategies_data:
        prompt += "\n\n## Strategy Information\n\n"
        # If a specific strategy is selected, only include that one
        if active_strategy and active_strategy in strategies_data:
            prompt += f"Selected Strategy: {active_strategy}\n\n"
            prompt += json.dumps({active_strategy: strategies_data[active_strategy]}, indent=2)
            prompt += ("\n\nIMPORTANT: You must ONLY discuss and recommend the selected strategy above. "
                       "Do not mention or suggest other strategies even if they might be relevant. "
                       "If asked about other strategies, politely redirect the conversation to focus on the "
                       "selected strategy or suggest clicking a different strategy button in the sidebar.")
        else:
            # Otherwise include all strategies
            prompt += json.dumps(strategies_data, indent=2)
    return prompt

# Sidebar for model and temperature selection
with st.sidebar:
   # st.markdown("<h1 style='text-align: center;'>Settings</h1>", unsafe_allow_html=True)
   # st.caption("Note: Gemini-1.5-pro-002 can only handle 2 requests per minute, gemini-1.5-flash-002 can handle 15 per minute")

    # Ensure model_name is initialized
    if 'model_name' not in st.session_state:
        st.session_state.model_name = "gemini-2.0-flash"  # default model

    #model_option = st.selectbox(
        #"Select Model:", ["gemini-2.0-pro-exp-02-05", "gemini-2.0-flash"]
    #)

    # Update model_name if it has changed
    #if model_option != st.session_state.model_name:
        #st.session_state.model_name = model_option
        #st.session_state.messages = []
        #st.session_state.chat_session = None

    # Add divider before strategy buttons
    st.divider()
    
    # Strategy section title
    st.markdown("<h1 style='text-align: center;'>Low-Intensity Strategies</h1>", unsafe_allow_html=True)
    
    # CSS styling for buttons
    custom_css = """
    <style>
    /* Target buttons within the sidebar */
    [data-testid="stSidebar"] button {
        width: 300px; /* Fixed width for all buttons */
        white-space: normal; /* Allow text to wrap */
        height: auto; /* Allow height to adjust based on content */
        background-color: #6A157D; /* Purple background */
        color: white; /* White text */
        border: none; /* Remove borders */
        padding: 10px 24px; /* Add some padding */
        text-align: center; /* Center text */
        text-decoration: none; /* Remove underline */
        display: inline-block; /* Get buttons to line up nicely */
        font-size: 10px; /* Increase font size */
        margin: 4px 2px; /* Add some margin */
        cursor: pointer; /* Add a pointer cursor on hover */
        border-radius: 12px; /* Rounded corners */
    }
    
    /* Add a hover effect for buttons */
    [data-testid="stSidebar"] button:hover {
        background-color: #871BA1; /* Darker purple */
        color: white !important; /* White text on hover - using !important to override any conflicting styles */
    }
    
    /* Style for active buttons - Updated selector to target disabled buttons */
    [data-testid="stSidebar"] button:disabled {
        background-color: #E1A2F0 !important; /* Light purple for active button */
        color: black !important; /* Make text visible against light background */
        opacity: 1 !important; /* Prevent the default opacity reduction */
        cursor: default;
    }
    
    /* Style for primary button (Return to Main Chat) */
    [data-testid="stSidebar"] [kind="primary"] {
        background-color: #C1E5F5 !important; /* Light blue background */
        color: black !important; /* Black text for better contrast */
        margin-bottom: 15px !important; /* Add space below */
    }
    </style>
    """

    # Set colors for sidebar  
    st.markdown(custom_css, unsafe_allow_html=True)
    
    # Wrap the strategy buttons in a container with a unique id so that only these buttons are affected
    st.markdown('<div id="strategy-buttons">', unsafe_allow_html=True)
    
    # Strategy buttons list
    strategies = [
        "Active Supervision",
        "Behavior-Specific Praise",
        "High-Probability Request Sequences",
        "Instructional Choice",
        "Instructional Feedback",
        "Opportunities to Respond",
        "Precorrection"
    ]
    
    # If a strategy is active, first display the return button using a primary button type
    if st.session_state.active_strategy:
        # Use type="primary" to make it visually different, styled with CSS
        if st.button("Return to Main Chat", key="return_button", type="primary", help="Go back to main chat"):
            st.session_state.active_strategy = None
            st.session_state.messages = []
            st.session_state.chat_session = None
            st.rerun()

    # Display strategy buttons
    for strategy in strategies:
        # Use a unique key for each button to avoid conflicts during re-renders
        button_key = f"strategy_button_{strategy}"
        # Determine if the current strategy is active
        is_active = st.session_state.active_strategy == strategy
        
        # Use disabled attribute for active strategy button which will be styled via CSS
        if is_active:
            st.button(strategy, key=button_key, disabled=True, help="Currently selected strategy")
        else:
            # Regular button for inactive strategies
            if st.button(strategy, key=button_key, help="Click to explore this strategy"):
                st.session_state.active_strategy = strategy  # Activate the strategy
                st.session_state.messages = []  # Clear chat history when switching strategies
                st.session_state.chat_session = None  # Reset session
                st.rerun()
    
    # Close the container for the strategy buttons
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Debug section - only include once in the sidebar
    st.divider()
    st.markdown("<h1 style='text-align: center;'>Debug Info</h1>", unsafe_allow_html=True)
    for debug_msg in st.session_state.debug:
        st.text(debug_msg)

# Create a main container for all content
main_container = st.container()

# Create a container for the funding acknowledgment that will appear at the bottom
funding_container = st.container()

# Now fill the main container with content
with main_container:
    # Title and BotDescription with dynamic header based on active strategy
    if st.session_state.active_strategy:
        st.markdown(f"<h2 style='background-color: #F0F2F6; padding: 15px;'>Focus on {st.session_state.active_strategy}</h2>", unsafe_allow_html=True)
        # First message intro for active strategy
        if not st.session_state.messages:
            strategy_intros = {
                "Active Supervision": "Active Supervision involves moving, scanning, and interacting with students to prevent and address behavior concerns.",
                "Behavior-Specific Praise": "Behavior-Specific Praise is a form of positive reinforcement that acknowledges specific student behaviors.",
                "High-Probability Request Sequences": "High-Probability Request Sequences involve making requests students are likely to comply with before making more challenging requests.",
                "Instructional Choice": "Instructional Choice involves embedding options into lessons for students to select based on their preferences.",
                "Instructional Feedback": "Instructional Feedback provides precise information to students about their academic, social, and behavioral performance.",
                "Opportunities to Respond": "Opportunities to Respond involves offering frequent opportunities for students to engage with academic material.",
                "Precorrection": "Precorrection involves proactively reminding students of expected behaviors before challenging situations arise."
            }
            intro = strategy_intros.get(st.session_state.active_strategy, "")
            st.markdown(f"<div style='background-color: #F0F2F6; padding: 15px;'>You're currently exploring the {st.session_state.active_strategy} strategy. {intro}</div>", unsafe_allow_html=True)
            st.markdown("<div style='background-color: #F0F2F6; padding: 15px;'>Ask questions about how to implement this strategy in your classroom or describe a scenario where you might use it.</div>", unsafe_allow_html=True)
        else:
            st.markdown(f"<div style='background-color: #F0F2F6; padding: 15px;'>You're currently exploring the {st.session_state.active_strategy} strategy. Ask questions about how to implement this strategy in your classroom or how it can help with specific scenarios.</div>", unsafe_allow_html=True)
    else:
        st.markdown("<h2 style='background-color: #F0F2F6; padding: 15px;'>Welcome to the Low-Intensity Strategies Bot!</h2>", unsafe_allow_html=True)
        st.markdown("<div style='background-color: #F0F2F6; padding: 15px;'>The goal of this bot is to assist you in selecting a low-intensity strategy that fits your needs—whether you are proactively planning for engagement in your lessons or responding to an interfering or challenging behavior you are experiencing.<br><br><strong>Directions:</strong> If you would like to explore multiple low-intensity strategy options, type a description of the scenario you are experiencing or a lesson plan idea into the chat to get started. If you would like to focus on one strategy specifically, click the name of the strategy on the side menu to get started.</div>", unsafe_allow_html=True) 
    st.caption("Note: This Bot is under development and can make mistakes. Visit ci3t.org for information and resources about low-intensity strategies.")
    
    # Add extra spacing between caption and chat input
    st.write("")
    st.write("")
    st.write("")

    # --- Add this block to initialize Firebase ---
    @st.cache_resource
    def initialize_firebase():
        try:
            # Construct the service account info dictionary from secrets
            # Streamlit secrets are accessed via st.secrets
            # Ensure the section name [firebase] matches your secrets.toml config
            firebase_secrets = {
                "type": st.secrets["firebase"]["type"],
                "project_id": st.secrets["firebase"]["project_id"],
                "private_key_id": st.secrets["firebase"]["private_key_id"],
                "private_key": st.secrets["firebase"]["private_key"],
                "client_email": st.secrets["firebase"]["client_email"],
                "client_id": st.secrets["firebase"]["client_id"],
                "auth_uri": st.secrets["firebase"]["auth_uri"],
                "token_uri": st.secrets["firebase"]["token_uri"],
                "auth_provider_x509_cert_url": st.secrets["firebase"]["auth_provider_x509_cert_url"],
                "client_x509_cert_url": st.secrets["firebase"]["client_x509_cert_url"],
                "universe_domain": st.secrets["firebase"]["universe_domain"]
            }

            # Check if firebase app is already initialized to prevent errors on rerun
            if not firebase_admin._apps:
                cred = credentials.Certificate(firebase_secrets)
                firebase_admin.initialize_app(cred)
                # You can add a debug message if needed, but avoid st.success here
                # st.session_state.debug.append("Firebase initialized successfully!")
            else:
                # App is already initialized, just get the existing app
                # st.session_state.debug.append("Firebase already initialized.")
                pass # Or firebase_admin.get_app() if you need the app object

            # Get Firestore client
            db = firestore.client()
            st.session_state.debug.append("Firestore client obtained.") # Debug message
            return db

        except Exception as e:
            st.error(f"Error initializing Firebase: {e}. Please check your secrets configuration.")
            st.session_state.debug.append(f"Firebase Init Error: {e}") # Debug message
            return None

    # Call the function to initialize Firebase and get the db client
    # This variable 'db' will hold the Firestore client object if initialization was successful
    db = initialize_firebase()
    
    # Initialize Gemini client
    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])

    # Display chat messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    # Add download buttons after the messages but before user input
    if st.session_state.messages:  # Only show if there are messages
        # Create a container with right-aligned content
        download_container = st.container()
        with download_container:
            # Adjusted column widths to make buttons shorter - using more columns to push content right
            _, _, _, format_col, button_col = st.columns([3, 3, 2, 2, 2])
            
            with format_col:
                format_option = st.selectbox(
                    "Format:",
                    ["Markdown (.md)", "PDF (.pdf)", "Word (.docx)"],
                    label_visibility="collapsed",
                    key="format_selection"
                )
            
            with button_col:
                # Determine file extension and data based on format selection
                if format_option == "PDF (.pdf)":
                    file_ext = ".pdf"
                    file_data = get_chat_pdf()
                    mime_type = "application/pdf"
                elif format_option == "Word (.docx)":
                    file_ext = ".docx"
                    file_data = get_chat_docx()
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                else:  # Default to Markdown
                    file_ext = ".md"
                    file_data = get_chat_text_markdown()
                    mime_type = "text/markdown"
                
                # Common base filename
                base_filename = f"LIS-Coach-Chat-{'strategy-' + st.session_state.active_strategy if st.session_state.active_strategy else 'main'}"
                
                st.download_button(
                    label="Download chat",
                    data=file_data,
                    file_name=f"{base_filename}{file_ext}",
                    mime=mime_type,
                    help="Save this conversation to your device",
                    key="download_chat",
                    use_container_width=True
                )

    # Handle form submission and generate response
    if st.session_state.should_generate_response:
        # Create combined prompt from responses
        combined_prompt = "Form Responses:\n"
        for q, a in st.session_state.form_responses.items():
            combined_prompt += f"{q}: {a}\n"
        
        # Add user message to chat history
        current_message = {"role": "user", "content": combined_prompt}
        st.session_state.messages.append(current_message)

        with st.chat_message("user"):
            st.markdown(current_message["content"])

        # Generate and display assistant response
        with st.chat_message("assistant"):
            message_placeholder = st.empty()

            # Initialize chat session if needed
            if st.session_state.chat_session is None:
                generation_config = {
                    "temperature": st.session_state.temperature,
                    "top_p": 0.95,
                    "top_k": 40,
                    "max_output_tokens": 8192,
                }
                model = genai.GenerativeModel(
                    model_name=st.session_state.model_name,
                    generation_config=generation_config,
                )
                
                # Build complete system prompt with active strategy if applicable
                complete_system_prompt = build_system_prompt(st.session_state.active_strategy)
                
                initial_messages = [
                    {"role": "user", "parts": [f"System: {complete_system_prompt}"]},
                    {"role": "model", "parts": ["Understood. I will follow these instructions."]},
                ]
                
                st.session_state.chat_session = model.start_chat(history=initial_messages)

            # Generate response with error handling
            try:
                response = st.session_state.chat_session.send_message(current_message["content"])
                full_response = response.text
                message_placeholder.markdown(full_response)
                st.session_state.messages.append({"role": "assistant", "content": full_response})
                st.session_state.debug.append("Assistant response generated")
                save_message(current_message["content"], full_response)
            except Exception as e:
                st.error(f"An error occurred while generating the response: {e}")
                st.session_state.debug.append(f"Error: {e}")

        st.session_state.should_generate_response = False
        st.rerun()

    # User input with context-aware placeholder
    placeholder_text = "Ask about how to use this strategy in your classroom" if st.session_state.active_strategy else "Describe a classroom scenario or ask about low-intensity strategies"
    user_input = st.chat_input(placeholder_text)

    if user_input:
        # Add user message to chat history
        current_message = {"role": "user", "content": user_input}
        st.session_state.messages.append(current_message)

        with st.chat_message("user"):
            st.markdown(current_message["content"])

        # Generate and display assistant response
        with st.chat_message("assistant"):
            message_placeholder = st.empty()

            # Prepare messages for Gemini API
            if st.session_state.chat_session is None:
                generation_config = {
                    "temperature": st.session_state.temperature,
                    "top_p": 0.95,
                    "top_k": 40,
                    "max_output_tokens": 8192,
                }
                model = genai.GenerativeModel(
                    model_name=st.session_state.model_name,
                    generation_config=generation_config,
                )
                
                # Build complete system prompt with active strategy if applicable
                complete_system_prompt = build_system_prompt(st.session_state.active_strategy)
                
                # Initialize chat with system prompt
                initial_messages = [
                    {"role": "user", "parts": [f"System: {complete_system_prompt}"]},
                    {"role": "model", "parts": ["Understood. I will follow these instructions."]},
                ]
                
                st.session_state.chat_session = model.start_chat(history=initial_messages)

            # Generate response with error handling
            try:
                response = st.session_state.chat_session.send_message(current_message["content"])
                full_response = response.text
                message_placeholder.markdown(full_response)
                st.session_state.messages.append({"role": "assistant", "content": full_response})
                st.session_state.debug.append("Assistant response generated")
                save_message(current_message["content"], full_response)
            except Exception as e:
                st.error(f"An error occurred while generating the response: {e}")
                st.session_state.debug.append(f"Error: {e}")

        st.rerun()

# Now put the funding acknowledgment in the funding container (will appear at the bottom)
with funding_container:
    st.markdown("<div style='text-align: center; margin-top: 20px;'><small style='color: rgb(128, 128, 128);'>This bot is programmed with information from ci3t.org.\n\nThis work was supported, in part, by ASU's Mary Lou Fulton Teachers College (MLFTC). The opinions and findings expressed in this document are those of the author and do not necessarily reflect those of the funding agency.</small></div>", unsafe_allow_html=True)
